import os
import requests
import fitz  # PyMuPDF
import json
import streamlit as st
import base64
from io import BytesIO
import pandas as pd
from datetime import datetime
import html
from docx import Document
from docx.shared import Pt, Inches

# --- PAGE CONFIGURATION ---
st.set_page_config(
    page_title="Tariff Impact Tracker",
    page_icon="📈",
    layout="wide"
)

# --- STYLING ---
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Poppins:wght@400;600;700&display=swap');
html, body, * {
    font-family: 'Poppins', sans-serif !important;
}
.block-container {
    padding-top: 4rem !important; padding-left: 3rem !important; padding-right: 3rem !important; padding-bottom: 3rem !important;
}
h1, h2, h3, .stTitle, .stHeader {
    font-family: 'Poppins', sans-serif !important; font-weight: 600 !important; color: #1e1e1e;
}
.aranca-header {
    display: flex; justify-content: space-between; align-items: center; padding-bottom: 1rem; border-bottom: 2px solid #f0f2f6; margin-bottom: 2rem;
}
.aranca-title {
    font-size: 2.0rem; font-weight: 700; color: #1e1e1e;
}
.aranca-logo img {
    height: 40px; object-fit: contain;
}
.report-card {
    background-color: #ffffff; border: 1px solid #e0e0e0; border-left: 5px solid #00416A; border-radius: 8px; padding: 20px; margin-bottom: 20px; box-shadow: 0 2px 4px rgba(0,0,0,0.05);
}
.report-card h3 {
    margin-top: 0; color: #00416A;
}
.report-card table {
    width: 100%; border-collapse: collapse;
}
.report-card th, .report-card td {
    padding: 10px 15px; text-align: left; border-bottom: 1px solid #e0e0e0; vertical-align: top;
}
.report-card th {
    background-color: #f9f9f9;
}
.report-card ul {
    padding-left: 20px; margin-top: 0;
}
</style>
""", unsafe_allow_html=True)


# --- API KEY & LOGO SETUP ---
try:
    DEEPSEEK_API_KEY = st.secrets.get("deepseek", {}).get("api_key")
    FMP_API_KEY = st.secrets.get("fmp", {}).get("api_key")
except (KeyError, FileNotFoundError):
    st.error("API keys (deepseek, fmp) not found in Streamlit secrets.")
    st.stop()

def get_base64_logo_image(path="logo.png"):
    if os.path.exists(path):
        with open(path, "rb") as f:
            return base64.b64encode(f.read()).decode()
    return ""

logo_base64 = get_base64_logo_image()

st.markdown(
    f"""
    <div class="aranca-header">
        <div class="aranca-title">Tariff Impact Tracker</div>
        <div class="aranca-logo">
            <img src="data:image/png;base64,{logo_base64}" alt="Aranca Logo">
        </div>
    </div>
    """,
    unsafe_allow_html=True
)

# --- CORE APPLICATION LOGIC (No changes to these functions) ---
@st.cache_data(ttl=3600)
def get_transcript_from_fmp(ticker, year, quarter):
    if not FMP_API_KEY:
        st.error("Error: FMP_API_KEY not found in secrets.")
        return None
    url = f"https://financialmodelingprep.com/api/v3/earning_call_transcript/{ticker}?quarter={quarter}&year={year}&apikey={FMP_API_KEY}"
    try:
        response = requests.get(url)
        response.raise_for_status()
        data = response.json()
        if data and "content" in data[0]:
            return data[0]["content"]
        else:
            st.warning(f"No transcript content found for {ticker} for Q{quarter} {year}.")
            return None
    except requests.exceptions.RequestException as e:
        st.error(f"Error fetching data from FMP API: {e}")
        return None
    except (IndexError, KeyError):
        st.error("Error parsing FMP API response. The data might be empty or in an unexpected format.")
        return None

def extract_text_from_pdf(uploaded_file):
    full_text = ""
    try:
        file_bytes = uploaded_file.getvalue()
        with fitz.open(stream=file_bytes, filetype="pdf") as doc:
            for page in doc:
                full_text += page.get_text() + "\n"
    except Exception as e:
        st.error(f"An error occurred while reading '{uploaded_file.name}': {e}")
    return full_text

@st.cache_data(ttl=3600)
def analyze_text_with_deepseek(text_content):
    if not DEEPSEEK_API_KEY:
        st.error("Error: DEEPSEEK_API_KEY not found in secrets.")
        return None
    if not text_content or not text_content.strip():
        st.warning("Input text is empty. Cannot perform analysis.")
        return None
    prompt = f"""
    As a specialized financial analyst, your task is to analyze the following corporate document.
    Your focus must be exclusively on comments related to **tariffs, trade duties, and import taxes**.
    **Critical Rule:** You must ignore all general financial metrics (e.g., overall revenue, total orders, EBITA) unless the text explicitly states that tariffs are the cause of the financial impact. If no specific financial data related to tariffs is mentioned, the corresponding fields in your response must be an empty list `[]`.
    Extract the information and structure your response as a valid JSON object. If a specific piece of information is not mentioned, use `null` or an empty list.
    JSON Specification:
    - **company_name**: The full company name mentioned in the document.
    - **quarterly_impact**: A list of objects detailing financial impacts **explicitly attributed to tariffs** in the current quarter. Examples: "Tariffs increased costs by $5M," or "Gross margin was impacted by 20 basis points due to import duties." If no such specific financial impact is mentioned, this MUST be an empty list.
    - **forward_guidance_impact**: A list of objects detailing future financial guidance **explicitly related to tariffs**. If none is mentioned, this MUST be an empty list.
    - **qualitative_impacts**: A list of strings describing non-financial impacts or general business environment effects due to tariffs. Examples: "Customer project delays due to tariff uncertainty," or "Increased complexity in supply chain planning."
    - **mitigation_strategies**: A list of strings detailing the specific strategies or actions the company is taking to handle the impact of tariffs.
    - **overall_sentiment**: Your assessment of the company's sentiment regarding tariffs ("Positive", "Neutral", "Negative"). This should be based only on the tariff-related comments.
    - **summary**: A brief, one-paragraph summary of the company's position on tariffs, synthesizing ONLY the specific findings. If the document provides few specifics, your summary must state that.
    Document Text:
    ---
    {text_content[:40000]}
    ---
    """
    headers = {"Content-Type": "application/json", "Authorization": f"Bearer {DEEPSEEK_API_KEY}"}
    data = {"model": "deepseek-chat", "messages": [{"role": "user", "content": prompt}], "temperature": 0.1, "response_format": {"type": "json_object"}}
    try:
        response = requests.post("https://api.deepseek.com/chat/completions", headers=headers, json=data, timeout=120)
        response.raise_for_status()
        content_str = response.json()['choices'][0]['message']['content']
        return json.loads(content_str)
    except requests.exceptions.RequestException as e:
        st.error(f"Error calling DeepSeek API: {e}")
        return None
    except (json.JSONDecodeError, KeyError) as e:
        st.error(f"Error parsing DeepSeek API JSON response: {e}")
        return None

# --- UI DISPLAY & FILE GENERATION FUNCTIONS (No changes in this section) ---
def display_tariff_report(company_name, analysis):
    """Displays the analysis for a single company using a standardized HTML table."""
    if not analysis:
        st.warning(f"No analysis data to display for {company_name}.")
        return

    st.header(f"Tariff Impact Analysis: {analysis.get('company_name', company_name)}")
    
    sentiment = analysis.get('overall_sentiment', 'N/A')
    summary = analysis.get('summary', 'No summary provided.')
    summary_html = f"""
    <div class="report-card">
        <h3>Executive Summary</h3>
        <p><strong>Overall Sentiment on Tariffs:</strong> {html.escape(sentiment)}</p>
        <p>{html.escape(summary)}</p>
    </div>
    """
    st.markdown(summary_html, unsafe_allow_html=True)

    def display_impact_table(title, impacts):
        if not impacts:
            st.markdown(f"""
            <div class="report-card">
                <h3>{title}</h3>
                <p><i>No specific financial impacts from tariffs were mentioned in the document.</i></p>
            </div>
            """, unsafe_allow_html=True)
            return
        if isinstance(impacts, dict):
            impacts = [impacts]
            
        df = pd.DataFrame(impacts)
        df_display = df.rename(columns={
            'metric': 'Metric',
            'impact_value': 'Impact',
            'unit': 'Unit',
            'source_quote': 'Source Quote'
        })
        df_display = df_display.fillna('NA')
        table_html = df_display.to_html(index=False, escape=False, border=0)
        full_html = f"""
        <div class="report-card">
            <h3>{title}</h3>
            {table_html}
        </div>
        """
        st.markdown(full_html, unsafe_allow_html=True)

    display_impact_table("Quarterly Financial Impact", analysis.get('quarterly_impact'))
    display_impact_table("Forward Guidance Impact", analysis.get('forward_guidance_impact'))

    # MODIFIED: Display Qualitative Impacts as a paragraph
    qualitative_impacts = analysis.get('qualitative_impacts')
    if qualitative_impacts:
        # Join list items into a single paragraph
        paragraph_text = " ".join(qualitative_impacts)
        impacts_html = f"""
        <div class="report-card">
            <h3>Qualitative Impacts</h3>
            <p>{html.escape(paragraph_text)}</p>
        </div>
        """
        st.markdown(impacts_html, unsafe_allow_html=True)

    # MODIFIED: Display Mitigation Strategies as a natural language sentence
    strategies = analysis.get('mitigation_strategies')
    if strategies:
        strategy_text = ""
        # Join the list of strategies into a flowing sentence
        if len(strategies) == 1:
            strategy_text = strategies[0]
        elif len(strategies) == 2:
            strategy_text = f"{strategies[0]} and {strategies[1]}"
        else:
            strategy_text = ", ".join(strategies[:-1]) + f", and {strategies[-1]}"
        
        full_paragraph = f"To manage these impacts, the company is employing several strategies, including {strategy_text}."
        
        strategies_html = f"""
        <div class="report-card">
            <h3>Mitigation Strategies</h3>
            <p>{html.escape(full_paragraph)}</p>
        </div>
        """
        st.markdown(strategies_html, unsafe_allow_html=True)

def create_comparison_table(all_analyses, period_source, year):
    st.header("Cross-Company Comparison")
    comparison_data = []
    for company_key, analysis in all_analyses.items():
        if not analysis: continue
        q_impacts = analysis.get('quarterly_impact', [])
        if isinstance(q_impacts, dict): q_impacts = [q_impacts]
        f_impacts = analysis.get('forward_guidance_impact', [])
        if isinstance(f_impacts, dict): f_impacts = [f_impacts]
        q_summary = "; ".join([f"{i.get('metric','N/A')}: {i.get('impact_value','N/A')}" for i in q_impacts]) or "Not specified"
        f_summary = "; ".join([f"{i.get('metric','N/A')}: {i.get('impact_value','N/A')}" for i in f_impacts]) or "Not specified"
        total_summary_parts = []
        if q_impacts: total_summary_parts.append(f"Q2 Impact: {q_summary}")
        if f_impacts: total_summary_parts.append(f"FY{year+1} Guidance: {f_summary}")
        total_summary = "<br>".join(total_summary_parts) or "No specific impact mentioned."
        mitigation_list = analysis.get('mitigation_strategies', [])
        mitigation_html = "<ul>" + "".join([f"<li>{s}</li>" for s in mitigation_list]) + "</ul>" if mitigation_list else "Not specified"
        comparison_data.append({"Company": f"<strong>{analysis.get('company_name', company_key)}</strong>", "Period / Source": period_source, "Tariff Impact Summary": total_summary, "Mitigation": mitigation_html})
    if not comparison_data:
        st.info("No data available for comparison.")
        return
    df = pd.DataFrame(comparison_data)
    table_html = df.to_html(index=False, escape=False, border=0)
    full_html = f"""<div class="report-card"><h3>Comparison Summary</h3>{table_html}</div>"""
    st.markdown(full_html, unsafe_allow_html=True)

def generate_html_report(all_analyses, period_source, year, logo_base64_string):
    styles = """
    <style>
        body { font-family: 'Poppins', sans-serif; }
        .report-card { background-color: #ffffff; border: 1px solid #e0e0e0; border-left: 5px solid #00416A; border-radius: 8px; padding: 20px; margin-bottom: 20px; box-shadow: 0 2px 4px rgba(0,0,0,0.05); }
        .report-card h3 { margin-top: 0; color: #00416A; }
        .report-card table { width: 100%; border-collapse: collapse; }
        .report-card th, .report-card td { padding: 10px 15px; text-align: left; border-bottom: 1px solid #e0e0e0; vertical-align: top; }
        .report-card th { background-color: #f9f9f9; }
        .report-card ul { padding-left: 20px; margin-top: 0; }
        .aranca-header { display: flex; justify-content: space-between; align-items: center; padding-bottom: 1rem; border-bottom: 2px solid #f0f2f6; margin-bottom: 2rem; }
        .aranca-title { font-size: 2.0rem; font-weight: 700; color: #1e1e1e; }
        .aranca-logo img { height: 40px; object-fit: contain; }
    </style>
    """
    header_html = f"""
    <div class="aranca-header">
        <div class="aranca-title">Tariff Impact Tracker</div>
        <div class="aranca-logo">
            <img src="data:image/png;base64,{logo_base64_string}" alt="Aranca Logo">
        </div>
    </div>
    """
    full_html_content = f"<html><head><title>Tariff Impact Report</title>{styles}</head><body>{header_html}"
    for company, analysis in all_analyses.items():
        if not analysis: continue
        full_html_content += f"<h2>Tariff Impact Analysis: {analysis.get('company_name', company)}</h2>"
        sentiment = analysis.get('overall_sentiment', 'N/A')
        summary = analysis.get('summary', 'No summary provided.')
        full_html_content += f"""<div class="report-card"><h3>Executive Summary</h3><p><strong>Overall Sentiment on Tariffs:</strong> {html.escape(sentiment)}</p><p>{html.escape(summary)}</p></div>"""
        for title, key in [("Quarterly Financial Impact", "quarterly_impact"), ("Forward Guidance Impact", "forward_guidance_impact")]:
            impacts = analysis.get(key)
            if not impacts:
                full_html_content += f"""<div class="report-card"><h3>{title}</h3><p><i>No specific financial impacts from tariffs were mentioned.</i></p></div>"""
            else:
                df = pd.DataFrame(impacts if isinstance(impacts, list) else [impacts]).fillna('NA')
                df_display = df.rename(columns={'metric': 'Metric', 'impact_value': 'Impact', 'unit': 'Unit', 'source_quote': 'Source Quote'})
                table_html = df_display.to_html(index=False, escape=False, border=0)
                full_html_content += f"""<div class="report-card"><h3>{title}</h3>{table_html}</div>"""
        qual_impacts = analysis.get('qualitative_impacts')
        if qual_impacts:
            full_html_content += '<div class="report-card"><h3>Qualitative Impacts</h3><ul>'
            for impact in qual_impacts:
                full_html_content += f"<li>{html.escape(impact)}</li>"
            full_html_content += "</ul></div>"
        strategies = analysis.get('mitigation_strategies')
        if strategies:
            full_html_content += '<div class="report-card"><h3>Mitigation Strategies</h3><ul>'
            for s in strategies:
                full_html_content += f"<li>{html.escape(s)}</li>"
            full_html_content += "</ul></div>"
        full_html_content += "<hr>"
    full_html_content += "</body></html>"
    return full_html_content

def generate_word_report(all_analyses, period_source, year):
    doc = Document()
    doc.add_heading('Tariff Impact Report', level=0)
    for company, analysis in all_analyses.items():
        if not analysis: continue
        doc.add_heading(f"Analysis for: {analysis.get('company_name', company)}", level=1)
        doc.add_heading('Executive Summary', level=2)
        doc.add_paragraph().add_run(f"Overall Sentiment on Tariffs: {analysis.get('overall_sentiment', 'N/A')}").bold = True
        doc.add_paragraph(analysis.get('summary', 'No summary provided.'))
        for title, key in [("Quarterly Financial Impact", "quarterly_impact"), ("Forward Guidance Impact", "forward_guidance_impact")]:
            doc.add_heading(title, level=2)
            impacts = analysis.get(key)
            if not impacts:
                doc.add_paragraph("No specific financial impacts from tariffs were mentioned.")
            else:
                df = pd.DataFrame(impacts if isinstance(impacts, list) else [impacts]).fillna('NA')
                df = df.rename(columns={'metric': 'Metric', 'impact_value': 'Impact', 'unit': 'Unit', 'source_quote': 'Source Quote'})
                table = doc.add_table(rows=1, cols=len(df.columns))
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                for i, col_name in enumerate(df.columns):
                    hdr_cells[i].text = col_name
                for index, row in df.iterrows():
                    row_cells = table.add_row().cells
                    for i, cell_value in enumerate(row):
                        row_cells[i].text = str(cell_value)
        qual_impacts = analysis.get('qualitative_impacts')
        if qual_impacts:
            doc.add_heading('Qualitative Impacts', level=2)
            for impact in qual_impacts:
                doc.add_paragraph(impact, style='List Bullet')
        strategies = analysis.get('mitigation_strategies')
        if strategies:
            doc.add_heading('Mitigation Strategies', level=2)
            for s in strategies:
                doc.add_paragraph(s, style='List Bullet')
        doc.add_page_break()
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# --- STREAMLIT UI LAYOUT ---
st.subheader("Data Source")
data_source = st.radio(
    "Choose where to get the transcript from:",
    ("Fetch from FMP API", "Upload PDF Transcript(s)"),
    horizontal=True,
    label_visibility="collapsed"
)

# MODIFIED: Initialize session state for storing results
if 'all_analysis_results' not in st.session_state:
    st.session_state.all_analysis_results = {}
if 'analysis_period' not in st.session_state:
    st.session_state.analysis_period = ""
if 'analysis_year' not in st.session_state:
    st.session_state.analysis_year = datetime.now().year

if data_source == "Fetch from FMP API":
    tickers_input = st.text_input("Company Ticker(s)", "AAPL, MSFT, GOOGL", help="Enter one or more tickers, separated by commas.")
    c2, c3 = st.columns(2)
    with c2:
        year = st.number_input("Year", min_value=2010, max_value=2030, value=2024)
    with c3:
        quarter = st.selectbox("Quarter", [1, 2, 3, 4], index=1)
    
    if st.button("Fetch & Analyze Transcripts", type="primary"):
        tickers = [ticker.strip().upper() for ticker in tickers_input.split(',') if ticker.strip()]
        if tickers:
            # MODIFIED: Clear previous results and update session state
            st.session_state.all_analysis_results = {}
            st.session_state.analysis_period = f"{year} Q{quarter} / Earnings Call"
            st.session_state.analysis_year = year
            
            with st.spinner("Generating analysis..."):
                for ticker in tickers:
                    text_to_analyze = get_transcript_from_fmp(ticker, year, quarter)
                    if text_to_analyze:
                        # MODIFIED: Store results in session state
                        st.session_state.all_analysis_results[ticker] = analyze_text_with_deepseek(text_to_analyze)

elif data_source == "Upload PDF Transcript(s)":
    uploaded_files = st.file_uploader("Upload one or more PDF files", type="pdf", accept_multiple_files=True)
    
    if st.button("Upload & Analyze PDFs", type="primary"):
        if uploaded_files:
            # MODIFIED: Clear previous results and update session state
            st.session_state.all_analysis_results = {}
            st.session_state.analysis_period = "Uploaded Docs"
            st.session_state.analysis_year = datetime.now().year
            
            with st.spinner("Generating analysis..."):
                for uploaded_file in uploaded_files:
                    company_name = os.path.splitext(uploaded_file.name)[0]
                    text_to_analyze = extract_text_from_pdf(uploaded_file)
                    if text_to_analyze:
                        # MODIFIED: Store results in session state
                        st.session_state.all_analysis_results[company_name] = analyze_text_with_deepseek(text_to_analyze)
        else:
            st.warning("Please upload at least one PDF file.")


# MODIFIED: Check session state for results, not a local variable
if st.session_state.all_analysis_results:
    st.markdown("---")
    
    # Display the results on the page
    for company, analysis in st.session_state.all_analysis_results.items():
        display_tariff_report(company, analysis)
        st.markdown("---")
    
    if len(st.session_state.all_analysis_results) > 1:
        create_comparison_table(st.session_state.all_analysis_results, st.session_state.analysis_period, st.session_state.analysis_year)
    
    st.markdown("---")
    st.header("Download Report")

    # Prepare content for download buttons
    html_content = generate_html_report(st.session_state.all_analysis_results, st.session_state.analysis_period, st.session_state.analysis_year, logo_base64)
    word_buffer = generate_word_report(st.session_state.all_analysis_results, st.session_state.analysis_period, st.session_state.analysis_year)

    # Create columns for download buttons
    col1, col2 = st.columns(2)

    with col1:
        st.download_button(
            label="📄 Download as HTML",
            data=html_content,
            file_name="tariff_impact_report.html",
            mime="text/html",
        )

    with col2:
        st.download_button(
            label="📄 Download as Word",
            data=word_buffer,
            file_name="tariff_impact_report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )